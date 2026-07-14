# -*- coding: utf-8 -*-
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_shading(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24') # 3pt width
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '4F46E5') # Indigo color border
    tcBorders.append(left)
    
    for side in ['top', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    
    # Add a blank paragraph with spacing after table
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(4)
    p_after.paragraph_format.space_after = Pt(4)

def format_run(run, font_name='Arial', size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4') # 0.5 pt
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'D1D5DB')
            tblBorders.append(border)
        tblPr[0].append(tblBorders)

def main():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Colors
    primary_color = (31, 41, 55) # charcoal
    accent_color = (79, 70, 229) # indigo
    body_color = (55, 65, 81) # dark gray
    heading_color = (30, 27, 75) # dark navy

    # --- Title Page ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(10)
    title_run = title_p.add_run("TÀI LIỆU PHÂN TÍCH CHI TIẾT TỪNG BƯỚC CỦA 11 LUỒNG HOẠT ĐỘNG DỰ ÁN MOTOV")
    format_run(title_run, 'Arial', 18, bold=True, color_rgb=accent_color)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(40)
    sub_run = sub_p.add_run("Tài liệu kỹ thuật đặc tả luồng hoạt động nghiệp vụ hệ thống (Web & Mobile)")
    format_run(sub_run, 'Arial', 12, italic=True, color_rgb=body_color)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(120)
    info_run = info_p.add_run("Hệ thống: Motov - Motorcycle Rental Platform\nThành phần đặc tả: Phân tích luồng xử lý Backend, Frontend & Database\nTài liệu tự động tạo lập chi tiết\nThời gian: 2026")
    format_run(info_run, 'Arial', 11, color_rgb=body_color)
    
    doc.add_page_break()

    # --- Section I: Overview ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    h1_run = h1.add_run("I. Tổng Quan Hệ Thống và Công Nghệ Sử Dụng")
    format_run(h1_run, 'Arial', 15, bold=True, color_rgb=heading_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(
        "Dự án Motov là nền tảng công nghệ hỗ trợ cho thuê xe máy trực tuyến. Hệ thống được tổ chức theo kiến trúc Client-Server, "
        "cho phép kết nối giữa Khách thuê xe, Chủ xe cung cấp phương tiện, Nhân viên điều phối và Quản trị viên hệ thống. "
        "Dự án bao gồm 3 phân hệ chính:"
    )
    format_run(run, 'Arial', 11, color_rgb=body_color)

    stack_details = [
        ("1. Frontend Web Client (client/):", " React (Vite), TypeScript, Tailwind CSS, Lucide React, Framer Motion."),
        ("2. Frontend Mobile App (mobile/):", " React Native (Expo), TypeScript, Redux Toolkit, React Navigation."),
        ("3. Backend Express Server (server/):", " Node.js + Express, TypeScript, TSX, MongoDB + Mongoose, Socket.io.")
    ]

    for title, desc in stack_details:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(title)
        format_run(run_title, 'Arial', 11, bold=True, color_rgb=primary_color)
        run_desc = p.add_run(desc)
        format_run(run_desc, 'Arial', 11, color_rgb=body_color)

    doc.add_page_break()

    # --- Section II: Detailed Flows ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    h1_run = h1.add_run("II. Đặc Tả Từng Bước 11 Luồng Hoạt Động Kỹ Thuật Chi Tiết")
    format_run(h1_run, 'Arial', 15, bold=True, color_rgb=heading_color)

    # LUỒNG 1
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("1. Luồng Xác Thực và Phân Quyền (Authentication & RBAC)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Khách hàng nhập Form đăng ký hoặc đăng nhập tại Client (Web/Mobile).\n"
        "Bước 2: Client gửi request POST đến /api/auth/register hoặc /api/auth/login kèm body (email, password).\n"
        "   - Trong trường hợp Google Auth: Client gửi mã idToken lên POST /api/auth/google.\n"
        "Bước 3: Server xử lý yêu cầu qua Router và Controller:\n"
        "   - Đăng ký (register): Mã hóa mật khẩu bằng bcrypt.hash() với độ muối (salt) = 10, tạo mới tài liệu trong collection 'users' với quyền mặc định roles = ['Customer'].\n"
        "   - Đăng nhập (login): Tìm email trong DB. So khớp mật khẩu đã lưu bằng bcrypt.compare(). Kiểm tra nếu trạng thái status === 'Suspended' (bị khóa) -> Trả về mã lỗi 403.\n"
        "   - Google Login: Gọi thư viện OAuth của Google để kiểm chứng token. Nếu hợp lệ, tìm hoặc tạo mới tài khoản người dùng.\n"
        "Bước 4: Tạo JWT Token:\n"
        "   - Server sử dụng jwt.sign() mã hóa payload (id, email, roles) với thời hạn 7 ngày dựa trên biến môi trường JWT_SECRET.\n"
        "Bước 5: Trả về Client:\n"
        "   - Response gửi mã JWT và thông tin user. Client lưu JWT vào localStorage (Web) hoặc SecureStore (Mobile).\n"
        "Bước 6: Gửi kèm và xác thực token ở các request sau:\n"
        "   - Mọi request cần bảo mật được Client đính kèm header 'Authorization: Bearer <token>'.\n"
        "   - Middleware 'authMiddleware' giải mã JWT, kiểm tra tính hợp lệ và lưu thông tin vào 'req.user'.\n"
        "   - Middleware 'restrictTo(...roles)' kiểm tra mảng req.user.roles để cấp quyền truy cập tài nguyên phù hợp."
    )

    # LUỒNG 2
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("2. Luồng Trở Thành Đối Tác/Chủ Xe (Become Owner)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Khách hàng (Customer) vào mục cá nhân và điền thông tin: số điện thoại ngân hàng, thông tin thanh toán (Tên ngân hàng, Số tài khoản, Chủ tài khoản) và tải lên hình chữ ký số cá nhân.\n"
        "Bước 2: Người dùng thực hiện ký thỏa thuận đối tác trực tuyến trên giao diện.\n"
        "Bước 3: Gửi request PUT đến /api/auth/become-owner kèm các thông số đăng ký và văn bản thỏa thuận.\n"
        "Bước 4: Server tiếp nhận qua hàm becomeOwner trong authController.ts:\n"
        "   - Cập nhật thông tin trong User collection: ownerRequestStatus = 'Pending', ownerContractSigned = true, ownerContractSignedAt = new Date(), ownerContractText = văn bản thỏa thuận.\n"
        "Bước 5: Tạo thông báo in-app Notification gửi tới toàn bộ tài khoản Admin/Staff thông báo có đối tác mới đăng ký.\n"
        "Bước 6: Phê duyệt từ phía Admin:\n"
        "   - Admin vào trang quản trị gọi API GET /api/auth/owner-requests để hiển thị các yêu cầu chờ duyệt.\n"
        "   - Admin duyệt yêu cầu qua API PUT /api/auth/owner-requests/:id/approve -> Cập nhật roles của user thêm 'Owner', đặt trạng thái duyệt thành 'Approved'.\n"
        "   - Admin từ chối qua API PUT /api/auth/owner-requests/:id/reject -> Đặt trạng thái 'Rejected' kèm lý do từ chối. Hệ thống tạo thông báo gửi tới người dùng."
    )

    # LUỒNG 3
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("3. Luồng Xác Thực Danh Tính Khách Hàng (eKYC / Identity Verification)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Người dùng truy cập trang cá nhân, thực hiện điền thông tin CCCD (Số CCCD, họ tên, ngày sinh, quê quán, địa chỉ thường trú) và tải lên 2 ảnh: Mặt trước & Mặt sau của thẻ CCCD.\n"
        "Bước 2: Giao diện Client gọi API upload ảnh lên server (POST /api/upload) -> Ảnh được lưu vào thư mục 'uploads/' trên đĩa cứng qua thư viện Multer và trả về link URL tĩnh.\n"
        "Bước 3: Client gửi request POST /api/auth/verify-identity chứa thông tin text và 2 link ảnh CCCD.\n"
        "Bước 4: Server tiếp nhận qua hàm submitIdentityVerification trong authController.ts:\n"
        "   - Lưu thông tin CCCD vào trường citizenIdInfo trong collection 'users'.\n"
        "   - Đặt trạng thái xác minh identityStatus = 'Pending'.\n"
        "   - Gửi thông báo Notification tới Admin và Staff thông báo có yêu cầu eKYC mới cần phê duyệt.\n"
        "Bước 5: Admin/Staff kiểm duyệt từ xa:\n"
        "   - Sử dụng API GET /api/auth/identity-requests để xem toàn bộ danh sách CCCD chờ duyệt.\n"
        "   - Nếu thông tin trùng khớp với ảnh: Gọi API PUT /api/auth/identity-requests/:id/approve -> Đổi trạng thái user thành identityStatus = 'Verified', đồng thời tạo thông báo in-app chúc mừng khách hàng.\n"
        "   - Nếu thông tin sai hoặc mờ: Gọi API PUT /api/auth/identity-requests/:id/reject -> Đổi trạng thái thành identityStatus = 'Rejected', kèm theo tin nhắn mô tả lý do từ chối để khách hàng thực hiện lại."
    )

    # LUỒNG 4
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("4. Luồng Quản Lý Xe Máy (Vehicle Management)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Chủ xe (Owner) hoặc Admin gửi yêu cầu đăng xe mới qua POST /api/vehicles với body chứa: tên mẫu xe, biển số, giá thuê, địa chỉ nhận xe, tọa độ vị trí [Lng, Lat], danh mục và hình ảnh.\n"
        "Bước 2: Server thực hiện tạo bản ghi trong collection 'vehicles':\n"
        "   - Đặt mặc định trạng thái status = 'Available'.\n"
        "   - Lưu thông tin gpsLocation dưới dạng GeoJSON Point: { type: 'Point', coordinates: [Lng, Lat] } để phục vụ định vị.\n"
        "   - Lưu liên kết ownerId đến ID của chủ xe đăng ký.\n"
        "Bước 3: Khách thuê tìm kiếm xe trên Client:\n"
        "   - Lấy danh sách toàn bộ xe máy hoạt động qua GET /api/vehicles.\n"
        "   - Tìm kiếm xe quanh vị trí hiện tại qua GET /api/vehicles/nearby: Server dùng toán tử địa lý $near của MongoDB để truy vấn các xe máy có khoảng cách gần nhất trong bán kính cho phép và sắp xếp theo khoảng cách tăng dần.\n"
        "   - Xem danh sách xe gợi ý qua GET /api/vehicles/recommendations.\n"
        "Bước 4: Khách hàng quản lý xe yêu thích:\n"
        "   - Khách có thể nhấn tim để thêm xe vào mục yêu thích qua API POST /api/users/favorites -> Hệ thống đẩy ID xe vào mảng favorites của User.\n"
        "Bước 5: Bảo trì xe máy:\n"
        "   - Chủ xe/Staff có thể gọi API PUT /api/vehicles/:id/maintenance-reset để ghi nhận bảo trì xe máy, reset cờ requiresMaintenance = false và cập nhật lastMaintenanceOdometer = odometer hiện tại."
    )

    # LUỒNG 5
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("5. Luồng Đặt Xe và Thanh Toán Đặt Cọc (Booking & VNPAY Payment)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Khách thuê điền thông tin đặt thuê trên Client và gửi yêu cầu POST /api/bookings.\n"
        "Bước 2: Server chạy hàm createBooking trong bookingController.ts:\n"
        "   - Kiểm tra xem user.identityStatus có phải là 'Verified' không -> Nếu không báo lỗi eKYC.\n"
        "   - Kiểm tra xe máy có rảnh hay không qua checkVehicleAvailability() bằng truy vấn tìm kiếm lịch trùng trong khoảng thời gian khách chọn.\n"
        "   - Tính số ngày thuê làm tròn theo mốc block 24 giờ. Tính tổng tiền thuê.\n"
        "   - Nếu có nhập promoCode, kiểm tra tính hợp lệ của Discount trong DB, trừ đi số tiền được chiết khấu và tăng số lượt sử dụng voucher.\n"
        "   - Tính toán số tiền: depositAmount = 30% tổng tiền; remainingAmount = 70% tổng tiền.\n"
        "   - Lưu đơn Booking với trạng thái status = 'Pending', ghi nhận snapshot thông tin xe vào đơn hàng.\n"
        "   - Tạo Notification in-app thông báo cho Khách thuê, Chủ xe và các Admin/Staff.\n"
        "   - Gửi Email xác nhận đã tạo đơn thành công cho khách hàng.\n"
        "Bước 3: Khách hàng tiến hành thanh toán trực tuyến tiền đặt cọc:\n"
        "   - Client gửi yêu cầu POST /api/bookings/:id/vnpay-url.\n"
        "   - Server tạo URL thanh toán VNPAY bằng cách băm thuật toán HMAC-SHA512 với mã bí mật vnp_HashSecret.\n"
        "   - Client chuyển hướng người dùng sang trang thanh toán của VNPAY.\n"
        "Bước 4: Xử lý Webhook IPN từ VNPAY:\n"
        "   - Khi giao dịch trên cổng VNPAY kết thúc, VNPAY gọi ngầm đến API /api/bookings/vnpay-ipn.\n"
        "   - Server xác thực chữ ký bảo mật, kiểm tra số tiền khớp với đơn cọc hay không.\n"
        "   - Nếu thanh toán thành công: Cập nhật booking.isPaid = true; booking.status = 'Confirmed'. Gửi thông báo đẩy chúc mừng thành công và lên lịch nhắc nhở tự động.\n"
        "   - Nếu thanh toán thất bại/bị hủy: Tự động đổi booking.status = 'Cancelled' và tạo thông báo hủy."
    )

    # LUỒNG 6
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("6. Luồng Vận Hành Bàn Giao và Nhận Lại Xe (Pickup, Return & Late Fees)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Bàn giao xe (Pick up):\n"
        "   - Khách thuê đến đại lý, Staff kiểm tra thông tin và gọi API PUT /api/bookings/staff/bookings/:id/pickup.\n"
        "   - Server lưu số km ban đầu của xe vào booking.startOdometer.\n"
        "   - Đổi booking.status = 'Ongoing' (đang đi) và cập nhật trạng thái xe máy trong DB thành status = 'Rented'.\n"
        "   - Kích hoạt lập lịch nhắc nhở trả xe tự động và gửi thông báo hành trình bắt đầu.\n"
        "Bước 2: Nhận lại xe (Return):\n"
        "   - Khách thuê trả lại xe máy. Staff nhập thời gian trả thực tế và số Odometer kết thúc, gửi request PUT /api/bookings/:id/return.\n"
        "Bước 3: Server xử lý tính toán phụ phí phạt tại returnMotorbike:\n"
        "   - So sánh thời điểm trả thực tế với thời điểm hẹn trả. Nếu trễ, tính phí phạt trễ hạn dựa trên số giờ trễ nhân với phí thuê/giờ. Thêm vào mảng surcharges bản ghi loại 'Late Return' và cộng vào tổng số tiền.\n"
        "   - Nếu trả xe sớm trước hẹn >= 2 tiếng: phạt mất hoàn toàn tiền cọc đã đóng trước, tính tiền thuê dựa trên số giờ thuê thực tế. Thêm phụ phí phạt tương ứng.\n"
        "   - Đổi booking.status = 'Completed', lưu số Odometer kết thúc vào đơn hàng.\n"
        "Bước 4: Cập nhật thông số kỹ thuật xe máy:\n"
        "   - Cập nhật số odometer hiện tại của xe máy bằng số Odometer lúc trả xe. Đổi trạng thái xe máy về 'Available'.\n"
        "   - Tính toán chu kỳ bảo dưỡng: odometer mới - lastMaintenanceOdometer. Nếu lớn hơn hoặc bằng maintenanceInterval -> Đổi cờ requiresMaintenance = true và gửi thông báo cảnh báo bảo dưỡng.\n"
        "Bước 5: Cảnh cáo từ Admin (nếu có vi phạm):\n"
        "   - Admin phản hồi lý do trả xe qua PUT /api/bookings/:id/return-response. Nếu tick cảnh cáo, hệ thống cộng User.strikes += 1. Nếu người dùng nhận đủ 3 strikes, tài khoản tự động bị khóa sang trạng thái 'Suspended'."
    )

    # LUỒNG 7
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("7. Luồng Lập Lịch và Nhắc Nhở Tự Động (Booking Reminder Scheduler)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Đăng ký lịch nhắc nhở tự động:\n"
        "   - Khi một Booking chuyển trạng thái, hệ thống gọi hàm handleBookingStatusTransitionReminders().\n"
        "   - Nếu chuyển sang 'Confirmed': Tạo 2 bản ghi nhắc nhở Nhận xe (24h_before_pickup và 2h_before_pickup) trong collection 'bookingreminders' với thời gian dự kiến gửi tương ứng.\n"
        "   - Nếu chuyển sang 'Ongoing': Tạo 2 bản ghi nhắc nhở Trả xe (24h_before_return và 2h_before_return).\n"
        "   - Nếu chuyển sang 'Cancelled' hoặc 'Completed': Cập nhật tất cả nhắc nhở còn lại của đơn hàng đó thành 'Cancelled' để dừng gửi.\n"
        "Bước 2: Chạy tiến trình ngầm (Cron-job):\n"
        "   - Hàm initBookingReminderScheduler() khởi động vòng lặp quét DB định kỳ mỗi 5 phút một lần.\n"
        "Bước 3: Quét tìm nhắc nhở đến hạn:\n"
        "   - Tiến trình tìm các bản ghi nhắc nhở có status === 'Pending', scheduledTime <= Giờ hiện tại, và retryCount < 3.\n"
        "Bước 4: Kiểm tra tính hợp lệ trước khi gửi:\n"
        "   - Xác minh trạng thái thực tế của Booking liên kết (Ví dụ: Đơn đặt xe phải chưa bị hủy, trạng thái phải đúng Confirmed đối với nhắc nhở nhận xe và đúng Ongoing đối với nhắc nhở trả xe). Nếu sai lệch, cập nhật nhắc nhở sang 'Cancelled'.\n"
        "Bước 5: Thực hiện gửi tin nhắn đa kênh:\n"
        "   - Email: Gọi sendPickupReminderEmail() / sendReturnReminderEmail() gửi thư điện tử.\n"
        "   - SMS: Gọi sendPickupReminderSms() / sendReturnReminderSms() gửi tin nhắn điện thoại.\n"
        "   - In-app: Tạo đối tượng Notification đẩy lên màn hình điện thoại/web của khách hàng.\n"
        "Bước 6: Ghi nhận kết quả:\n"
        "   - Cập nhật trạng thái nhắc nhở thành 'Sent' và lưu sentTime. Nếu lỗi, tăng retryCount += 1 (Tối đa 3 lần thử lại)."
    )

    # LUỒNG 8
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("8. Luồng Trò Chuyện Trực Tuyến Thời Gian Thực (Real-time Chat)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Người dùng (Khách thuê hoặc Chủ xe) truy cập trang trò chuyện.\n"
        "Bước 2: Thiết lập kết nối WebSocket thông qua thư viện Socket.io:\n"
        "   - Client gửi request kết nối kèm JWT token.\n"
        "   - Hàm initSocket() trong socket.ts xác thực tính hợp lệ của token, giải mã lấy userId và lưu vào thông tin kết nối.\n"
        "   - Socket tự động tham gia vào phòng cá nhân có tên 'user_${userId}'.\n"
        "Bước 3: Người dùng bắt đầu hoặc tiếp tục cuộc hội thoại:\n"
        "   - Client lấy danh sách cuộc trò chuyện cũ qua GET /api/chats/conversations.\n"
        "   - Khi mở phòng chat, client gửi sự kiện 'join_conversation' kèm ID cuộc hội thoại để gia nhập phòng chat ảo.\n"
        "Bước 4: Thực hiện gửi và phân phối tin nhắn:\n"
        "   - Client gửi tin nhắn mới lên REST API: POST /api/chats/messages.\n"
        "   - Server lưu trữ tin nhắn vào collection 'messages', cập nhật trường lastMessageAt trong collection 'conversations'.\n"
        "   - Server gọi hàm getIO() để lấy đối tượng Socket.io, phát sự kiện 'new_message' chứa nội dung tin nhắn trực tiếp đến phòng cuộc trò chuyện tương ứng.\n"
        "Bước 5: Đồng bộ giao diện phía người nhận:\n"
        "   - Trình duyệt phía đối phương đang mở màn hình chat (trong cùng phòng conversationId) sẽ bắt lấy sự kiện và vẽ tin nhắn mới lên giao diện tức thời."
    )

    # LUỒNG 9
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("9. Luồng Thông Báo Đẩy Thời Gian Thực (Server-Sent Events - SSE)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Client đăng nhập và mở một kết nối HTTP dài hạn (Server-Sent Events) tới endpoint /api/notifications/stream.\n"
        "Bước 2: Server thiết lập và giữ kết nối mở thông qua hàm streamNotifications trong notificationController.ts:\n"
        "   - Trả về Header: Content-Type: text/event-stream và Connection: keep-alive.\n"
        "   - Đăng ký kết nối này vào danh sách quản lý activeClients trong bộ nhớ RAM qua hàm registerClient(userId, res).\n"
        "Bước 3: Duy trì kết nối ổn định:\n"
        "   - Cứ mỗi 30 giây, một tiến trình ngầm gửi chuỗi ': ping' tới toàn bộ kết nối đang mở để tránh bị các hệ thống mạng tự động đóng kết nối.\n"
        "Bước 4: Phát tán thông báo tự động (Mongoose Hook):\n"
        "   - Khi hệ thống lưu một bản ghi thông báo mới vào cơ sở dữ liệu qua Notification.create(...).\n"
        "   - Trigger hook post('save') trên Schema Notification tự động bắt sự kiện lưu dữ liệu thành công.\n"
        "   - Hook tiến hành import động realtimeService.ts và gọi hàm sendRealtimeNotification(userId, notificationDoc).\n"
        "   - Dịch vụ lọc các kết nối Response trong activeClients khớp với userId tương ứng và thực hiện ghi chuỗi JSON thông báo qua kết nối: res.write('data: ' + JSON.stringify(data)).\n"
        "Bước 5: Client đón nhận và hiển thị:\n"
        "   - Giao diện client bắt sự kiện eventSource.onmessage và hiển thị thông báo pop-up đẹp mắt ngay lập tức."
    )

    # LUỒNG 10
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("10. Luồng Giám Sát IoT Pin và Hành Trình Xe (Battery Monitoring & GPS)")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "Bước 1: Thiết bị IoT thông minh gắn trên xe máy định kỳ (ví dụ: mỗi 1 phút) gửi gói tin chứa dữ liệu dung lượng pin (batteryLevel), số km đã chạy (odometer), và tọa độ địa lý thực tế [Lng, Lat] lên API của Server.\n"
        "Bước 2: Server tiếp nhận dữ liệu và lưu trữ lịch sử vào collection 'batterylogs'.\n"
        "   - Collection được cấu hình dạng MongoDB Timeseries collection (với timeField là 'timestamp' và metaField là 'vehicleId') giúp tối ưu dung lượng lưu trữ và tốc độ truy vấn đối với dữ liệu chuỗi thời gian.\n"
        "   - Cấu hình chỉ mục không gian địa lý 2dsphere trên trường gpsLocation để tối ưu các truy vấn bản đồ.\n"
        "Bước 3: Hiển thị hành trình trực quan trên giao diện Client:\n"
        "   - Khi khách hàng hoặc chủ xe xem bản đồ giám sát (BikesMap), client gọi API lấy các bản ghi log pin và hành trình mới nhất của xe máy.\n"
        "   - Giao diện sử dụng bản đồ tích hợp (Leaflet/Google Maps) để vẽ lại lộ trình và hiển thị pin xe theo thời gian thực."
    )

    # LUỒNG 11
    h2 = doc.add_heading(level=2)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)
    h2_run = h2.add_run("11. Luồng Cấu Hình Hệ Thống và Seeder Dữ Liệu")
    format_run(h2_run, 'Arial', 12.5, bold=True, color_rgb=accent_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Mô tả hoạt động chi tiết:")
    format_run(run, 'Arial', 11, bold=True, color_rgb=primary_color)

    add_code_block(doc,
        "A. Thay đổi cấu hình động của hệ thống:\n"
        "   1. Admin truy cập trang cấu hình hệ thống, sửa đổi các tham số nhạy cảm như: Tài khoản kết nối sandbox VNPAY (vnp_TmnCode, vnp_HashSecret, vnp_Url), thông tin mail server SMTP, hoặc biểu phí phạt trễ hạn.\n"
        "   2. Client gửi request PUT /api/system/settings cập nhật các tham số.\n"
        "   3. Server lưu các tham số này trực tiếp vào collection 'systemsettings' dưới dạng các cặp key-value.\n"
        "   4. Khi chạy các nghiệp vụ cần thiết (như tạo link thanh toán VNPAY, gửi email nhắc nhở), backend thực hiện truy vấn động trong DB thông qua hàm helper getSettingVal() thay vì đọc từ file .env cứng. Điều này giúp hệ thống hoạt động linh hoạt, thay đổi cấu hình tức thời mà không cần khởi động lại máy chủ.\n"
        "\n"
        "B. Seeder dữ liệu mẫu cho môi trường kiểm thử (Development Environment):\n"
        "   1. Khi khởi động dự án hoặc chạy script cài đặt, hàm seedUsers(), seedFeedbacks(), seedChats() được kích hoạt.\n"
        "   2. Hệ thống kiểm tra xem dữ liệu đã có sẵn trong cơ sở dữ liệu chưa. Nếu chưa có, tự động tạo lập:\n"
        "      - Tạo 4 tài khoản mẫu đại diện cho 4 vai trò: Admin, Staff, Owner, Customer với mật khẩu được băm sẵn để nhà phát triển đăng nhập nhanh.\n"
        "      - Tạo các bản ghi đánh giá xe mẫu và tin nhắn chat mẫu giúp giao diện hiển thị dữ liệu chân thực ngay từ đầu."
    )

    doc.add_page_break()

    # --- Section III ---
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    h1_run = h1.add_run("III. Tổng Hợp Database Collections và Schema Fields")
    format_run(h1_run, 'Arial', 15, bold=True, color_rgb=heading_color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(
        "Bảng tổng hợp cấu trúc các Collection quan trọng trong MongoDB phục vụ vận hành hệ thống cho thuê xe máy Motov:"
    )
    format_run(run, 'Arial', 11, color_rgb=body_color)

    # Add Table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    headers = ["Collection Name", "Model File", "Key Schema Fields", "Purpose & Notes"]
    for i, title in enumerate(headers):
        hdr_cells[i].paragraphs[0].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(title)
        format_run(run, 'Arial', 10, bold=True, color_rgb=(255,255,255))
        set_cell_shading(hdr_cells[i], "4F46E5") # Indigo header background
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)

    db_info = [
        ("users", "User.ts", "username, email, passwordHash, roles (Customer/Owner/Staff/Admin), identityStatus, strikes, usedVouchers", "Quản lý thông tin tài khoản, trạng thái khóa, eKYC, tích hợp hợp đồng và voucher."),
        ("vehicles", "Vehicle.ts", "vehicleModel, licensePlate, rentalPrice, ownerId, status (Available/Rented/Maintenance), odometer, requiresMaintenance", "Lưu trữ thông tin xe máy, biển số, định giá thuê, quãng đường và trạng thái bảo trì."),
        ("bookings", "Booking.ts", "bookingCode, userId, vehicleId, pickupDateTime, returnDateTime, totalAmount, depositAmount, status, surcharges, endOdometer", "Ghi nhận toàn bộ thông tin đặt xe, trạng thái hành trình, thanh toán cọc và phí phạt."),
        ("bookingreminders", "BookingReminder.ts", "bookingId, reminderType, scheduledTime, status (Pending/Sent/Cancelled), channel, retryCount", "Lập lịch gửi email và tin nhắn SMS tự động nhắc nhở khách hàng trước giờ nhận/trả xe."),
        ("payments", "Payment.ts", "bookingId, userId, amount, paymentMethod (VNPAY/Cash), paymentStatus, transactionId, paymentDate", "Lưu trữ lịch sử giao dịch thanh toán đặt cọc hoặc thanh toán hoàn tất đơn hàng."),
        ("notifications", "Notification.ts", "userId, title, message, type (BookingPending/System), relatedId, isRead", "Lưu thông tin thông báo, tích hợp trigger tự động đẩy dữ liệu thời gian thực qua kết nối SSE."),
        ("discounts", "Discount.ts", "voucherCode, discountType, discountValue, usageLimit, usedCount, startDate, endDate", "Quản lý các chương trình ưu đãi, mã voucher giảm giá, giới hạn lượt dùng hệ thống."),
        ("conversations", "Conversation.ts", "participants (User IDs), lastMessageAt, unreadCount", "Quản lý các cuộc trò chuyện giữa khách hàng và chủ xe hoặc bộ phận hỗ trợ."),
        ("messages", "Message.ts", "conversationId, senderId, messageText, isRead, createdAt", "Lưu chi tiết nội dung tin nhắn chat thời gian thực phục vụ đồng bộ Socket."),
        ("batterylogs", "BatteryLog.ts", "vehicleId, timestamp, batteryLevel, odometer, gpsLocation (Point coordinates)", "Timeseries Collection lưu trữ dữ liệu định vị GPS và dung lượng pin gửi lên từ thiết bị IoT của xe máy."),
        ("systemsettings", "SystemSetting.ts", "key, value, description, updatedAt", "Cấu hình động các thông số hệ thống như tài khoản VNPAY, phí trễ hạn, máy chủ SMTP.")
    ]

    for row_idx, (col_name, model_name, fields, purpose) in enumerate(db_info):
        row_cells = table.add_row().cells
        data_row = [col_name, model_name, fields, purpose]
        bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
        for i, text in enumerate(data_row):
            row_cells[i].paragraphs[0].text = ""
            p = row_cells[i].paragraphs[0]
            run = p.add_run(text)
            
            # Formatting
            if i == 0:
                format_run(run, 'Consolas', 9.5, bold=True, color_rgb=primary_color)
            elif i == 1:
                format_run(run, 'Consolas', 9.5, color_rgb=accent_color)
            elif i == 2:
                format_run(run, 'Arial', 9.5, color_rgb=(75,85,99))
            else:
                format_run(run, 'Arial', 9.5, color_rgb=body_color)
                
            set_cell_shading(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=100, right=100)

    # Save to file
    output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '../TaiLieu_PhanTich_LuongHoatDong_Motov_ChiTiet.docx'))
    doc.save(output_path)
    print(f"SUCCESS: Document generated successfully at {output_path}")

if __name__ == '__main__':
    main()
